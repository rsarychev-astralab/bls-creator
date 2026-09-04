import io
import re
from io import BytesIO

import httpx
from docx import Document
from openpyxl import load_workbook
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2.service_account import Credentials

from app.config import NOTION_TOKEN, sa_path
from app.sources.notion import NOTION_VERSION, notion_page_id
from app.sources.sheet import SCOPES

_DRIVE_ID_RE = [
    re.compile(r"[?&]id=([a-zA-Z0-9_-]+)"),
    re.compile(r"/d/([a-zA-Z0-9_-]+)"),
    re.compile(r"/file/d/([a-zA-Z0-9_-]+)"),
    re.compile(r"/open\?id=([a-zA-Z0-9_-]+)"),
]


def drive_file_id(url: str) -> str:
    for pat in _DRIVE_ID_RE:
        m = pat.search(url or "")
        if m:
            return m.group(1)
    return ""


def _drive():
    creds = Credentials.from_service_account_file(str(sa_path()), scopes=SCOPES)
    return build("drive", "v3", credentials=creds, cache_discovery=False)


def xlsx_to_text(data: bytes) -> str:
    wb = load_workbook(BytesIO(data), data_only=True, read_only=True)
    parts = []
    try:
        for sheet in wb.worksheets:
            if sheet.title:
                parts.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
    finally:
        wb.close()
    return "\n".join(parts).strip()


def text_from_upload(name: str, data: bytes, mime: str = "") -> str:
    if not data:
        return ""
    lower = (name or "").lower()
    mime = (mime or "").lower()
    if lower.endswith(".xlsx") or "spreadsheetml.sheet" in mime:
        return xlsx_to_text(data)
    if lower.endswith(".docx") or "wordprocessingml" in mime:
        try:
            return docx_to_text(data)
        except Exception:
            return ""
    if data[:2] == b"PK":
        try:
            return xlsx_to_text(data)
        except Exception:
            try:
                return docx_to_text(data)
            except Exception:
                return ""
    return data.decode("utf-8-sig", errors="replace").strip()


def docx_to_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
            if any(cells):
                parts.append(" || ".join(cells))
    return "\n".join(parts).strip()


def fetch_questionnaire_from_notion(crm_url: str) -> dict:
    page_id = notion_page_id(crm_url)
    if not page_id:
        return {"ok": False, "error": "Нет id страницы CRM", "text": "", "source": "notion"}
    if not NOTION_TOKEN:
        return {"ok": False, "error": "Пустой NOTION_TOKEN", "text": "", "source": "notion"}
    headers = {
        "Authorization": f"Bearer {NOTION_TOKEN}",
        "Notion-Version": NOTION_VERSION,
    }
    try:
        with httpx.Client(timeout=30.0, verify=False, headers=headers) as client:
            resp = client.get(f"https://api.notion.com/v1/pages/{page_id}")
            data = resp.json()
            if resp.status_code != 200:
                return {
                    "ok": False,
                    "error": data.get("message") or f"Notion {resp.status_code}",
                    "text": "",
                    "source": "notion",
                }
            files = ((data.get("properties") or {}).get("Анкета БЛС") or {}).get("files") or []
        if not files:
            return {
                "ok": False,
                "error": "В Notion пустое поле «Анкета БЛС»",
                "text": "",
                "source": "notion",
            }
        texts = []
        names = []
        errors = []
        for item in files:
            name = item.get("name") or ""
            names.append(name)
            ftype = item.get("type") or "file"
            url = (item.get(ftype) or {}).get("url") or ""
            if "docs.google.com" in name or "drive.google.com" in name:
                drive = fetch_questionnaire(name)
                if drive.get("ok") and drive.get("text"):
                    texts.append(drive["text"])
                else:
                    errors.append(drive.get("error") or f"Drive: {name}")
                continue
            if not url:
                errors.append(f"Нет url у файла {name}")
                continue
            # S3 signed URL ломается, если тащить Notion Authorization
            with httpx.Client(timeout=30.0, verify=False) as raw:
                downloaded = raw.get(url)
            if downloaded.status_code != 200:
                errors.append(f"{name}: download {downloaded.status_code}")
                continue
            ctype = downloaded.headers.get("content-type", "")
            if (
                "spreadsheetml.sheet" in ctype
                or name.lower().endswith(".xlsx")
                or "wordprocessingml" in ctype
                or name.lower().endswith(".docx")
            ):
                texts.append(text_from_upload(name, downloaded.content, ctype))
            else:
                texts.append(downloaded.text)
        text = "\n\n".join(t for t in texts if t).strip()
        if not text:
            return {
                "ok": False,
                "error": "; ".join(errors) or "Не удалось прочитать анкету из Notion",
                "name": ", ".join(names),
                "text": "",
                "source": "notion",
            }
        return {
            "ok": True,
            "error": "; ".join(errors),
            "name": ", ".join(names),
            "text": text,
            "source": "notion",
        }
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "error": str(exc), "text": "", "source": "notion"}


def fetch_questionnaire(url: str) -> dict:
    file_id = drive_file_id(url)
    if not file_id:
        return {"ok": False, "error": "Нет ссылки на опросник", "text": ""}
    try:
        service = _drive()
        meta = (
            service.files()
            .get(fileId=file_id, fields="id,name,mimeType")
            .execute()
        )
        mime = meta.get("mimeType", "")
        text = ""
        if mime.startswith("application/vnd.google-apps."):
            export_mime = "text/plain"
            if mime == "application/vnd.google-apps.spreadsheet":
                export_mime = "text/csv"
            data = service.files().export(fileId=file_id, mimeType=export_mime).execute()
            text = data.decode("utf-8", errors="replace") if isinstance(data, bytes) else str(data)
        else:
            buf = io.BytesIO()
            req = service.files().get_media(fileId=file_id)
            downloader = MediaIoBaseDownload(buf, req)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            raw = buf.getvalue()
            text = text_from_upload(meta.get("name", ""), raw, mime)
        return {
            "ok": True,
            "error": "",
            "id": file_id,
            "name": meta.get("name", ""),
            "mime_type": mime,
            "text": text.strip(),
        }
    except HttpError as exc:
        return {
            "ok": False,
            "error": f"Drive: {exc.status_code} {exc.reason}. Расшарь файл на сервисный аккаунт.",
            "id": file_id,
            "text": "",
        }
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "error": str(exc), "id": file_id, "text": ""}
